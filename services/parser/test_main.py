import io
import unittest
from unittest.mock import patch

from fastapi import HTTPException
from docx import Document
from pypdf import PdfWriter

from main import (
    ParseRequest,
    decode_text_bytes,
    parse_docx_document,
    parse_document,
    parse_pdf_document,
    parse_text_document,
)


class ParserMainTestCase(unittest.TestCase):
    def test_decode_text_bytes_supports_gb18030(self):
        payload = "项目范围".encode("gb18030")
        self.assertEqual(decode_text_bytes(payload), "项目范围")

    def test_parse_text_document_recognizes_markdown_heading(self):
        result = parse_text_document("# 项目范围\n\n第8节 发布说明".encode("utf-8"))

        self.assertEqual(result["page_count"], 1)
        self.assertEqual(result["blocks"][0]["block_type"], "heading")
        self.assertEqual(result["blocks"][0]["text"], "项目范围")
        self.assertEqual(result["blocks"][0]["metadata_json"]["locator"]["line_start"], 1)

    def test_parse_text_document_keeps_document_line_ranges(self):
        result = parse_text_document("第一段第一行\n第一段第二行\n\n第二段".encode("utf-8"))

        self.assertEqual(result["blocks"][0]["metadata_json"]["locator"]["line_start"], 1)
        self.assertEqual(result["blocks"][0]["metadata_json"]["locator"]["line_end"], 2)
        self.assertEqual(result["blocks"][1]["metadata_json"]["locator"]["line_start"], 4)

    def test_parse_docx_document_preserves_heading_styles(self):
        document = Document()
        document.add_heading("发布手册", level=1)
        document.add_heading("上线范围", level=2)
        document.add_paragraph("上线前需完成回归测试。")
        buffer = io.BytesIO()
        document.save(buffer)

        result = parse_docx_document(buffer.getvalue())

        self.assertEqual(result["blocks"][0]["block_type"], "heading")
        self.assertEqual(result["blocks"][0]["text"], "发布手册")
        self.assertEqual(result["blocks"][1]["heading_path"], ["发布手册", "上线范围"])
        self.assertEqual(result["blocks"][2]["heading_path"], ["发布手册", "上线范围"])

    def test_parse_docx_document_extracts_tables_in_order(self):
        document = Document()
        document.add_heading("发布计划", level=1)
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "阶段"
        table.cell(0, 1).text = "负责人"
        table.cell(1, 0).text = "灰度发布"
        table.cell(1, 1).text = "平台团队"
        document.add_paragraph("灰度发布前应完成回归测试和通知。")
        buffer = io.BytesIO()
        document.save(buffer)

        result = parse_docx_document(buffer.getvalue())

        self.assertEqual(result["blocks"][1]["block_type"], "table")
        self.assertIn("阶段 | 负责人", result["blocks"][1]["text"])
        self.assertEqual(result["blocks"][1]["heading_path"], ["发布计划"])
        self.assertEqual(result["blocks"][2]["heading_path"], ["发布计划"])

    def test_parse_pdf_document_rejects_blank_text_pdf_without_images(self):
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=300)
        buffer = io.BytesIO()
        writer.write(buffer)

        with self.assertRaises(HTTPException) as context, patch(
            "main.extract_pdf_page_texts",
            return_value=[
                {
                    "page_no": 1,
                    "width": 300.0,
                    "height": 300.0,
                    "text": "",
                    "has_image": False,
                }
            ],
        ):
            parse_pdf_document(buffer.getvalue())

        self.assertEqual(context.exception.status_code, 422)
        self.assertEqual(context.exception.detail["code"], "pdf_no_extractable_content")
        self.assertEqual(context.exception.detail["ocr_provider"], None)

    def test_parse_pdf_document_explicitly_reports_disabled_ocr_when_image_pages_need_it(self):
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=300)
        buffer = io.BytesIO()
        writer.write(buffer)

        with self.assertRaises(HTTPException) as context, patch(
            "main.extract_pdf_page_texts",
            return_value=[
                {
                    "page_no": 1,
                    "width": 300.0,
                    "height": 300.0,
                    "text": "",
                    "has_image": True,
                }
            ],
        ), patch.dict(
            "os.environ",
            {
                "PARSER_OCR_PROVIDER": "disabled",
            },
            clear=False,
        ):
            parse_pdf_document(buffer.getvalue())

        self.assertEqual(context.exception.status_code, 422)
        self.assertEqual(context.exception.detail["code"], "ocr_required")
        self.assertEqual(context.exception.detail["ocr_provider"], "disabled")

    def test_parse_pdf_document_uses_mock_ocr_provider_for_blank_pdf(self):
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=300)
        buffer = io.BytesIO()
        writer.write(buffer)

        with patch(
            "main.extract_pdf_page_texts",
            return_value=[
                {
                    "page_no": 1,
                    "width": 300.0,
                    "height": 300.0,
                    "text": "",
                    "has_image": True,
                }
            ],
        ), patch.dict(
            "os.environ",
            {
                "PARSER_OCR_PROVIDER": "mock",
                "PARSER_OCR_MOCK_TEXT": "# 发布总览\n\n第8节 上线检查\n\n发布前需完成回归测试。",
            },
            clear=False,
        ):
            result = parse_pdf_document(buffer.getvalue())

        self.assertEqual(result["source"]["mode"], "ocr")
        self.assertEqual(result["source"]["ocr_provider"], "mock")
        self.assertEqual(result["blocks"][0]["block_type"], "heading")
        self.assertEqual(result["blocks"][1]["heading_path"], ["发布总览", "第8节 上线检查"])
        self.assertEqual(result["blocks"][2]["heading_path"], ["发布总览", "第8节 上线检查"])
        self.assertEqual(result["blocks"][1]["metadata_json"]["locator"]["page_line_start"], 3)

    def test_parse_pdf_document_only_ocrs_pages_that_lack_native_text(self):
        with patch(
            "main.extract_pdf_page_texts",
            return_value=[
                {
                    "page_no": 1,
                    "width": 612.0,
                    "height": 792.0,
                    "text": "原生文本页",
                    "has_image": True,
                },
                {
                    "page_no": 2,
                    "width": 612.0,
                    "height": 792.0,
                    "text": "",
                    "has_image": True,
                },
            ],
        ), patch("main.get_ocr_provider") as get_ocr_provider:
            provider = get_ocr_provider.return_value
            provider.name = "mock"
            provider.extract_pdf_pages.return_value = [
                {
                    "page_no": 2,
                    "width": 612.0,
                    "height": 792.0,
                    "text": "OCR 补回的第二页",
                }
            ]

            result = parse_pdf_document(b"%PDF-mock%")

        provider.extract_pdf_pages.assert_called_once_with(b"%PDF-mock%", page_numbers=[2])
        self.assertEqual(result["source"]["mode"], "ocr")
        self.assertEqual(result["source"]["ocr_provider"], "mock")
        self.assertEqual([page["text_length"] for page in result["pages"]], [5, 10])

    def test_parse_document_reads_storage_and_dispatches_by_logical_path(self):
        request = ParseRequest(
            workspace_id="ws_123",
            document_version_id="dv_123",
            storage_key="workspaces/ws_123/uploads/source.bin",
            sha256="abc123",
            logical_path="资料库/notes.md",
        )

        with patch("main.get_object_bytes", return_value="# 项目范围".encode("utf-8")) as get_object_bytes:
            with patch("main.parse_text_document", return_value={"ok": True}) as parse_text:
                result = parse_document(request)

        get_object_bytes.assert_called_once_with("workspaces/ws_123/uploads/source.bin")
        parse_text.assert_called_once_with("# 项目范围".encode("utf-8"))
        self.assertEqual(result, {"ok": True})

    def test_parse_request_allows_library_only_payloads(self):
        request = ParseRequest(
            library_id="lib_123",
            document_version_id="dv_123",
            storage_key="libraries/lib_123/source.bin",
            sha256="abc123",
        )

        self.assertIsNone(request.workspace_id)
        self.assertEqual(request.library_id, "lib_123")


if __name__ == "__main__":
    unittest.main()
